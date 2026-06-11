from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "半句项目概述与技术文档_详细版.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x28, 0x38)
MUTED = RGBColor(0x66, 0x72, 0x85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_east_asia(run, font_name: str = FONT_CN) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def style_cell_text(cell, bold=False, color=INK, size=10) -> None:
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        for run in p.runs:
            set_east_asia(run)
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_east_asia(run)
        run.bold = bold
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_east_asia(r)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_east_asia(r)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_east_asia(r)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_east_asia(r2)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], LIGHT_GRAY)
        style_cell_text(hdr[i], bold=True, color=DARK_BLUE, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            style_cell_text(cells[i], size=9.3)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x2A, 0x34, 0x44)
    doc.add_paragraph()


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = FONT_LATIN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = FONT_LATIN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("半句项目说明书")
    set_east_asia(r)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("基于短剧剧情理解的即时互动激发系统")
    set_east_asia(r)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    return doc


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("半句：基于短剧剧情理解的即时互动激发系统")
    set_east_asia(r)
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("项目概述、技术文档与总结自评")
    set_east_asia(r)
    r.font.size = Pt(15)
    r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    add_callout(
        doc,
        "一句话概述",
        "半句是一个面向短剧高光时刻的 AI 互动产品：系统理解剧情、识别高光、治理弹幕、生成片尾二创和声音资产，并在观看端以低门槛互动、贴图动效、同看社交和陪看战报形成完整体验闭环。",
    )
    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [1.6, 4.9])
    items = [
        ("文档版本", "V1.0 交付说明版"),
        ("更新时间", "2026-06-11"),
        ("当前主线", "电脑端 Web 展示版；Android 原生客户端为并行探索支线"),
        ("适用场景", "比赛提交、答辩讲解、项目验收、后续团队接手"),
    ]
    for row, (k, v) in zip(meta.rows, items):
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        style_cell_text(row.cells[0], bold=True, color=DARK_BLUE)
        style_cell_text(row.cells[1])
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    doc.add_heading("一、项目概述", level=1)
    doc.add_heading("1.1 项目思路", level=2)
    add_paragraph(
        doc,
        "短剧的情绪爆点通常发生在几十秒内，用户想表达“爽、震惊、心疼、好笑、磕到了”等情绪，但传统弹幕和评论需要打字，会打断观看节奏。半句的核心思路是：用 AI 先理解短剧内容，把高光时刻、情绪类型、互动文案和贴图策略结构化，然后由服务端下发给观看端，在正确时间点触发轻量互动，让用户用一次点击、一次选择或短暂狂点完成情绪表达。",
    )
    add_callout(
        doc,
        "项目定位",
        "本项目不是单纯的视频播放器，也不是简单弹幕系统，而是“大模型内容理解 + 人工复核 + 服务端策略下发 + Web 端沉浸渲染 + 用户反馈回流”的短剧互动闭环。",
    )
    doc.add_heading("1.1.1 Web 主线与 Android 支线关系", level=3)
    add_paragraph(
        doc,
        "当前项目采用“双线并行、主次明确”的推进方式。Web 端是当前可交付主线，承担比赛展示、功能验证、视觉打磨、复核后台和文档交付；Android 端是原生迁移支线，用于验证未来移动 App 的播放器、全屏、系统权限和原生交互能力。两条线共享同一套服务端接口、内容数据、高光时间轴、弹幕治理结果、二创资产和声音缓存，但当前交付质量以 Web 主线为准。",
    )
    add_table(
        doc,
        ["线路", "当前定位", "承担内容", "判断标准"],
        [
            ["Web 主线", "当前主要交付和演示版本。", "首页、播放页、高光互动、弹幕、贴图、片尾二创、声音播放、小半陪看、同看社交、复核后台和文档说明。", "必须稳定、好看、可录屏、可答辩。"],
            ["Android 支线", "原生客户端迁移探索，不影响 Web 主线。", "验证原生播放器、全屏、触摸、系统权限、后续移动端交互和 App 包装可能性。", "先跑通最小闭环，再逐步追赶 Web 体验。"],
            ["共享服务端", "两条客户端线共用的产品底座。", "FastAPI 接口、SQLite/未来 PostgreSQL、视频/封面/高光/弹幕/二创/声音资产。", "接口语义稳定，避免为单端写死。"],
        ],
        [1.25, 1.65, 2.65, 0.95],
    )

    doc.add_heading("1.2 核心功能介绍", level=2)
    add_table(
        doc,
        ["模块", "已实现能力", "与课题关系"],
        [
            [
                "短剧内容理解",
                "已支持短剧、剧集、封面和视频导入；高光点按时间轴存储；20 集完成 LLM 高光表达与贴图策略升级；复核页可调整高光名称、时间、来源、模型版本和贴图时间窗。",
                "对应“理解短剧内容、识别剧情高光点”。",
            ],
            [
                "高光互动",
                "播放页监听视频进度，到达高光点后触发选择题、情绪按钮、疯狂点击、贴图动效和群体比例反馈；互动结果上报服务端。",
                "对应“端上即时互动激发”。",
            ],
            [
                "弹幕治理",
                "支持轻聊、狂欢、沉浸三档弹幕；导入弹幕数据后经过规则、时间感知、语义分类、聚类去重、小模型、大模型复审候选、人工复核七层治理。",
                "解决弹幕低俗、剧透、刷屏和遮挡剧情的问题。",
            ],
            [
                "AIGC 片尾二创",
                "北往第一集已形成片尾剧情预测和图片分镜流程：用户选择主分支，再在分镜节点选择饮料、回家方式、车辆类型等个性化选项；图片和语音采用预生成缓存策略。",
                "对应“AIGC 内容生成”和短剧结尾延展体验。",
            ],
            [
                "声音资产",
                "个人页支持上传/录入授权声音样本，后端保存 voice profile，为预设台词生成 mp3 并缓存；片尾二创可选择原版声音或用户声音带入版。",
                "形成用户代入感，但避免实时自由生成带来的风险。",
            ],
            [
                "小半陪看",
                "固定 AI 陪看员“小半”在高光前后做轻提示，并在片尾生成观看战报，展示互动选择、点击表现、奖励和称号。",
                "强化陪伴感和演示记忆点。",
            ],
            [
                "同看与社交",
                "已实现好友申请、聊天、同看房间、房间动态、称号/徽章展示、逛逛动态流等社交雏形。",
                "让互动从单人情绪表达扩展到共同观看和分享。",
            ],
        ],
        [1.35, 3.65, 1.5],
    )

    doc.add_heading("1.3 技术实现与部署环境", level=2)
    add_table(
        doc,
        ["层级", "当前实现", "后续生产化方向"],
        [
            ["客户端/前端", "电脑端 Web 是当前主线，包含登录、选片首页、播放页、复核页、我的、聊聊、逛逛等页面；Android 原生端为并行支线。", "Web 主线继续打磨；Android 支线逐步接入同一套服务端接口，后续再评估 iOS。"],
            ["服务端", "FastAPI 提供短剧、剧集、高光、弹幕、同看、社交、声音和二创接口。", "云服务器部署，Nginx 反向代理，HTTPS 正式访问。"],
            ["数据库", "SQLite 本地数据库，适合演示和快速迭代。", "正式上线前迁移 PostgreSQL。"],
            ["文件资产", "本地保存视频、封面、贴图、二创图片、音频缓存和头像。", "迁移对象存储，并对声音样本、生成音频做权限控制。"],
            ["AI 能力", "大模型离线生成高光表达、贴图策略、二创文案；声音服务生成 mp3；小模型用于弹幕治理雏形。", "保持离线生成 + 缓存 + 人工复核，避免实时成本和风险。"],
            ["公网演示", "可使用临时公网隧道展示。", "后续使用域名、HTTPS、备份、监控和回滚方案。"],
        ],
        [1.35, 3.0, 2.15],
    )

    doc.add_heading("1.4 产品亮点与创新点", level=2)
    add_bullets(
        doc,
        [
            "剧情时间轴驱动互动：互动不是固定弹窗，而是与短剧情绪峰值、台词和剧情揭晓时间绑定。",
            "低门槛表达：用户无需输入长评论，只需点击、选择、狂点或点赞弹幕，即可完成情绪表达。",
            "贴图和动效与剧情联动：冲突、反转、爽点、甜蜜、虐心、搞笑等高光采用不同贴图、点击反馈和消失逻辑。",
            "七层弹幕治理：不仅判断弹幕内容，还结合剧情时间处理剧透，使用聚类降低审核量，并预留小模型和大模型复审。",
            "片尾 AI 二创从“看完就结束”变为“看完可延展”：用户选择分支后进入图片分镜和语音播放，形成可展示的 AIGC 亮点。",
            "用户声音带入：用授权声音样本生成预设台词，既增强代入感，又避免开放式声音滥用。",
            "小半陪看和成长体系：通过陪看战报、积分、称号、徽章展馆和同看房间增强长期留存。",
        ]
    )

    doc.add_heading("1.5 技术关键点", level=2)
    add_bullets(
        doc,
        [
            "AI 侧：大模型主要放在离线链路，用于高光表达、贴图策略、弹幕复审候选和二创提示词；小模型先用于弹幕治理的低成本判断。",
            "服务端侧：FastAPI 统一管理内容、互动、弹幕、社交、声音和二创数据，并保留 source、model_version、confidence 等可追踪字段。",
            "前端侧：播放页监听时间轴，在不遮挡剧情的前提下叠加高光弹层、贴图、弹幕、二创入口和全屏适配。",
            "部署侧：当前本地 Web 演示稳定优先；未来按单应用服务器、PostgreSQL、对象存储、HTTPS、备份、监控逐步迁移。",
        ]
    )


def add_tech_selection_expansion(doc: Document) -> None:
    doc.add_heading("2.3.1 选型决策方法", level=3)
    add_paragraph(
        doc,
        "本项目的技术选型不是按“技术先进程度”排序，而是按当前项目阶段判断：比赛展示需要先证明完整闭环，后续上线再逐步增强工程能力。因此选型优先级为：体验稳定性、开发速度、AI 链路可解释性、数据可回溯、后续迁移成本。暂时不为未来假设过度设计。",
    )
    add_table(
        doc,
        ["判断维度", "具体含义", "对本项目的影响"],
        [
            ["体验稳定性", "用户观看时不能被模型延迟、接口失败、弹层遮挡和播放卡顿影响。", "AI 生成尽量离线完成，端上只加载缓存结果。"],
            ["开发速度", "当前需求变化快，产品体验还在打磨，技术栈必须便于快速改动。", "Web + FastAPI + SQLite 更适合当前迭代。"],
            ["可解释性", "答辩时需要说明高光为什么触发、弹幕为什么被拦截、二创内容从哪里来。", "保留 source、model_version、confidence、review_status 等字段。"],
            ["可迁移性", "演示期不能过早引入复杂部署，但未来需要能迁移到正式环境。", "先本地文件和 SQLite，文档中规划 PostgreSQL 和对象存储。"],
        ],
        [1.4, 2.55, 2.55],
    )

    doc.add_heading("2.3.2 客户端 / 前端选型展开", level=3)
    add_table(
        doc,
        ["候选方案", "优点", "不足", "本项目判断"],
        [
            ["原生 Android", "播放器、全屏、权限和系统能力更强，适合最终 App。", "开发周期长，UI 达到当前 Web 质感成本高；当前原生探索效果未达标。", "暂不作为主线。"],
            ["原生 iOS", "体验质量上限高，适合精致客户端。", "开发环境和真机调试要求高，比赛迭代速度不如 Web。", "后续再评估。"],
            ["Flutter / React Native", "跨端能力强，后续可同时覆盖 Android/iOS。", "视频播放器、覆盖层、弹幕、全屏和原生权限仍需大量适配。", "不是当前最快路径。"],
            ["Web", "开发快，前后端联调成本低，适合电脑端展示和临时公网演示。", "手机端体验和原生能力有限。", "当前选择。"],
        ],
        [1.25, 1.75, 2.25, 1.25],
    )
    add_paragraph(
        doc,
        "最终选择电脑端 Web 作为主线，是因为当前阶段最关键的是把观看体验、高光触发、弹幕、片尾二创、声音播放和后台复核稳定串起来。Web 能最快验证产品逻辑，也方便在侧边栏和浏览器中持续调试。Android 原生端不是被放弃，而是作为支线推进：它先复用服务端接口做最小闭环，等 Web 主线体验稳定后，再逐步补齐播放器外层视觉、高光弹层、弹幕轨道、二创覆盖层和系统权限能力。",
    )
    add_table(
        doc,
        ["能力", "Web 主线状态", "Android 支线状态", "共用/差异"],
        [
            ["短剧列表和登录", "已作为主线体验入口。", "可复用 API 做原生页面。", "数据共用，UI 分端实现。"],
            ["视频播放", "当前展示质量以 Web 为准，支持控制栏、全屏和覆盖层。", "支线先接入原生视频播放，再逐步追赶 Web 交互。", "视频地址共用，播放器实现不同。"],
            ["高光互动", "已支持多高光时间轴、弹层、贴图、点击和上报。", "先做时间触发和上报，再做复杂动效。", "高光数据共用，渲染层分端实现。"],
            ["弹幕", "已支持三档模式、点击、点赞、回复和治理结果展示。", "支线后续接入弹幕轨道。", "弹幕治理共用，渲染性能需原生单独优化。"],
            ["片尾二创和声音", "已形成全屏分镜、选项和音频播放。", "支线第二批接入。", "图片、台词和音频缓存共用，页面交互分端实现。"],
        ],
        [1.25, 2.0, 2.0, 1.25],
    )

    doc.add_heading("2.3.3 服务端选型展开", level=3)
    add_table(
        doc,
        ["候选方案", "适合场景", "本项目取舍"],
        [
            ["FastAPI", "Python 生态、AI 脚本、数据处理、接口自动文档、快速迭代。", "当前选择；能和模型标注、弹幕治理、声音服务脚本自然衔接。"],
            ["Spring Boot", "大型企业后端、复杂权限、强类型工程治理。", "稳定但开发和配置成本更高，不适合当前快速探索期。"],
            ["Node.js", "前后端同语言、实时通信生态成熟。", "适合同看实时消息，但 AI 数据处理仍要回到 Python；当前不是主服务首选。"],
        ],
        [1.25, 2.55, 2.7],
    )
    add_paragraph(
        doc,
        "最终选择 FastAPI，是因为本项目 AI 链路较重：大模型标注、弹幕治理、小模型训练、声音服务和素材脚本都更贴近 Python 生态。FastAPI 可以同时满足接口开发速度、自动 API 文档和后续部署需求。真正进入生产后，可在 FastAPI 外层增加 Nginx、进程守护、日志、监控和任务队列，而不需要马上重写服务端。",
    )

    doc.add_heading("2.3.4 数据库与缓存选型展开", level=3)
    add_table(
        doc,
        ["候选方案", "优点", "不足", "当前结论"],
        [
            ["SQLite", "零部署、可复制、可回滚，适合演示和快速迭代。", "并发、权限、备份和远程访问能力有限。", "当前使用。"],
            ["PostgreSQL", "生产能力强，支持 JSON、索引、事务、备份和多人访问。", "需要部署、迁移和运维。", "正式上线前迁移。"],
            ["MySQL", "常见、成熟、学习资料多。", "JSON 和复杂结构表达不如 PostgreSQL 灵活。", "不是当前优先。"],
            ["Redis", "适合同看房间状态、在线成员、实时消息和热点缓存。", "不是持久数据库，需要额外维护。", "后续社交同看增强时引入。"],
        ],
        [1.2, 1.9, 2.0, 1.4],
    )
    add_paragraph(
        doc,
        "当前使用 SQLite 的理由是项目还在频繁调整数据结构，且主要用于本地演示和比赛提交。高光、弹幕、二创、声音资产等表结构仍在变化，如果现在就迁移 PostgreSQL，会增加迁移脚本和运维负担。未来当重点剧集体验稳定、数据结构基本固定、需要多人访问时，再把 SQLite 迁移到 PostgreSQL，并把同看临时状态交给 Redis。",
    )

    doc.add_heading("2.3.5 文件资产与对象存储选型展开", level=3)
    add_table(
        doc,
        ["资产类型", "当前处理", "未来处理", "原因"],
        [
            ["视频和封面", "本地视频库和本地封面路径。", "对象存储 + CDN 或签名 URL。", "视频文件大，本地路径不适合公网长期访问。"],
            ["贴图和徽章", "前端静态资源目录。", "对象存储或随前端静态资源发布。", "多数是公共资源，权限要求低。"],
            ["二创图片", "预生成后缓存到本地。", "对象存储，数据库保存 key。", "需要稳定加载和便于复核。"],
            ["声音样本", "本地私有文件。", "私有对象存储，严格权限控制。", "属于高风险隐私资产，不能公开访问。"],
            ["生成音频", "mp3 缓存文件。", "对象存储，按权限或签名地址播放。", "避免重复生成，降低成本。"],
        ],
        [1.25, 1.65, 1.75, 1.85],
    )
    add_paragraph(
        doc,
        "当前保留本地文件系统，是为了方便调试视频、贴图、二创图片和语音缓存。正式上线时必须迁移对象存储，因为浏览器公网访问需要稳定 URL，服务器磁盘也不能长期承载大量视频和音频。数据库只保存对象 key、hash、大小、归属和权限，文件本体交给对象存储管理。",
    )

    doc.add_heading("2.3.6 AI 模型与生成策略选型展开", level=3)
    add_table(
        doc,
        ["能力", "考虑过的方案", "最终策略", "原因"],
        [
            ["高光识别", "纯人工、规则关键词、大模型、小模型。", "大模型草稿 + 人工复核；小模型后置。", "短期数据量不足，人工复核能保证展示质量。"],
            ["弹幕审核", "全人工、规则、实时大模型、分层治理。", "七层治理。", "真实弹幕量大，单一人工或实时大模型都不可持续。"],
            ["片尾二创", "实时视频生成、实时图片生成、预生成图片分镜。", "预生成图片分镜 + 缓存语音。", "视频生成成本和稳定性不足，图片分镜更可控。"],
            ["声音带入", "浏览器端生成、服务端实时生成、预设文本缓存。", "服务端生成 mp3 并缓存。", "保护模型文件和用户声音样本，避免每次重复生成。"],
        ],
        [1.25, 2.0, 1.8, 1.45],
    )
    add_paragraph(
        doc,
        "AI 策略的核心是“不要把不稳定模型放到用户观看主链路”。高光、贴图、二创和声音都采用离线或半离线生成，复核后缓存成稳定数据。用户观看时，客户端只负责加载服务端已经确认过的时间轴、贴图、图片和音频。这样即使模型服务临时不可用，也不会影响基础观看和互动。",
    )

    doc.add_heading("2.3.7 端上渲染和互动实现选型展开", level=3)
    add_table(
        doc,
        ["实现点", "候选方式", "当前方案", "选择理由"],
        [
            ["视频播放", "第三方播放器、原生 video、自研播放器。", "HTML5 video + 自定义控制层。", "便于控制高光触发、全屏覆盖和二创入口。"],
            ["弹幕渲染", "Canvas、DOM、第三方弹幕库。", "DOM/CSS 轨道。", "弹幕量用于演示足够，点击点赞回复更容易实现。"],
            ["贴图动效", "视频内合成、Canvas、DOM/CSS 动画。", "DOM/CSS 贴图层。", "易按时间窗控制位置、点击区域和消失动画。"],
            ["片尾二创", "跳转新页面、弹窗、全屏覆盖层。", "全屏沉浸覆盖层。", "不打断播放页上下文，也能适配全屏观看。"],
        ],
        [1.25, 1.75, 1.65, 1.85],
    )
    add_paragraph(
        doc,
        "端上渲染最终选择“HTML5 video + 自定义覆盖层”，是因为本项目的核心不是单纯播放，而是在播放时间轴上叠加高光弹层、贴图、弹幕、二创入口和小半提示。自定义覆盖层能精确控制出现时机、安全区、点击反馈和全屏适配。Canvas 或第三方播放器在复杂交互上反而会增加修改成本。",
    )

    doc.add_heading("2.3.8 部署路径选型展开", level=3)
    add_table(
        doc,
        ["阶段", "可选方式", "当前选择", "原因"],
        [
            ["本地开发", "本地服务、Docker、云开发机。", "本地 FastAPI + SQLite。", "调试视频和本地模型最方便。"],
            ["比赛演示", "本地局域网、临时公网隧道、正式域名。", "本地 + 临时公网隧道。", "能快速给外部查看，不提前承担生产运维。"],
            ["准生产", "单机部署、容器编排、云函数。", "单应用服务器 + Nginx。", "当前规模不需要复杂微服务。"],
            ["正式上线", "SQLite、本机文件、PostgreSQL、对象存储。", "PostgreSQL + 对象存储 + HTTPS。", "保证数据安全、资源加载和备份恢复。"],
        ],
        [1.15, 1.8, 1.75, 1.8],
    )
    add_paragraph(
        doc,
        "部署路径不立即进入生产，是因为当前最有价值的工作仍是 Web 主线观看体验、重点剧集样板、内容复核和 AI 资产质量。Android 支线的存在不会改变生产化顺序：无论未来是 Web、公网 H5、Android 原生还是 iOS，底层都应该先稳定服务端接口、数据结构、声音授权和内容治理边界，再进入 PostgreSQL、对象存储、HTTPS、备份和监控迁移。这个顺序可以避免为了某个端的部署而牺牲产品迭代速度。",
    )


def add_feature_deep_dive(doc: Document) -> None:
    doc.add_heading("2.5 特色模块深度展开：弹幕与声音", level=2)
    add_paragraph(
        doc,
        "本项目最能体现“短剧即时互动激发”的不是单个按钮，而是围绕观看情绪建立的配套体验。弹幕模块解决的是“用户想表达但不能破坏观看”的问题；声音模块解决的是“用户希望带入剧情但不能实时等待模型和牺牲隐私”的问题。这两个模块都采用了分层设计，而不是简单堆功能。",
    )

    doc.add_heading("2.5.1 弹幕模块：从普通评论到时间感知互动层", level=3)
    add_paragraph(
        doc,
        "弹幕在短剧里有双重价值：一方面能制造“大家一起看”的热闹感，另一方面也最容易破坏体验，例如遮挡关键画面、提前剧透、刷屏、低俗内容和无关内容。因此本项目没有把弹幕当成普通文本列表，而是把它作为一个需要治理、分发和交互的观看层。",
    )
    add_table(
        doc,
        ["设计目标", "具体做法", "为什么重要"],
        [
            ["不打断观看", "提供轻聊、狂欢、沉浸三档模式，并支持字号、透明度、速度和显示区域调整。", "不同用户对弹幕接受度不同，不能强迫所有用户看同一密度弹幕。"],
            ["避免剧透", "弹幕审核结合剧情时间轴和关键揭晓时间，例如结局前不能提前出现“摩托车回家”等信息。", "短剧反转和悬念依赖揭晓节奏，剧透会直接破坏高光。"],
            ["保留氛围", "狂欢模式允许适度重复“哈哈哈”“救命”“磕到了”等情绪弹幕。", "重复弹幕本身就是短视频社区情绪共振的一部分，不应一刀切删除。"],
            ["可互动", "弹幕支持点击、点赞、回复，后续可携带用户头像、昵称、称号和好友关系。", "弹幕不只是飘过的文字，也可以成为同看社交入口。"],
            ["可复核", "后台可查看治理原因、风险标签、适合模式和复核状态。", "评委和后续运营需要知道为什么某条弹幕被放行、延后或拦截。"],
        ],
        [1.3, 3.0, 2.2],
    )
    add_paragraph(
        doc,
        "弹幕模块最关键的取舍是：不能只追求热闹，也不能为了安全把弹幕全部关掉。最终方案是把弹幕拆成“内容治理 + 密度模式 + 端上交互 + 后台复核”四个层次，让用户、运营和模型各自承担合适的职责。",
    )
    add_code_block(
        doc,
        [
            "弹幕 CSV / 人工导入",
            "  -> 文本清洗、颜文字和表情归类",
            "  -> 绑定剧集与建议出现时间",
            "  -> 七层治理：规则、时间感知、语义、聚类、小模型、LLM 候选、人工复核",
            "  -> 分发到轻聊 / 狂欢 / 沉浸模式",
            "  -> 播放页渲染、点击、点赞、回复",
            "  -> 用户行为回传，反向优化治理策略",
        ],
    )
    doc.add_page_break()
    add_table(
        doc,
        ["治理层", "处理内容", "技术考虑", "当前实现价值"],
        [
            ["规则层", "脏话、广告、联系方式、过长文本、明显刷屏。", "最快、最便宜，适合确定性问题。", "先拦截明显风险，减少后续模型压力。"],
            ["时间感知层", "根据剧情揭晓时间判断是否剧透。", "弹幕不能只看文本，还要看出现在哪一秒。", "解决短剧悬念和反转最怕提前透露的问题。"],
            ["语义分类层", "判断相关性、低俗、出戏、情绪类型、适合模式。", "可由大模型离线批量审核，不建议实时逐条调用。", "让弹幕分类更贴近剧情语义。"],
            ["聚类去重层", "把大量相似弹幕归为一组。", "真实弹幕会大量重复，先聚类再审核代表项。", "降低人工和模型审核成本。"],
            ["小模型层", "输出 pass_score、confidence、model_version。", "用已有审核结果训练低成本判断器。", "为后续规模化弹幕治理做准备。"],
            ["LLM 复审候选层", "低置信度、高风险、高价值弹幕进入队列。", "大模型只处理值得复审的样本。", "控制成本，同时保留复杂语义判断能力。"],
            ["人工复核层", "高赞弹幕、推荐池内容、模型不确定样本。", "人工只看少数关键内容。", "保证最终展示内容可控、可解释。"],
        ],
        [1.15, 1.75, 2.0, 1.6],
    )
    add_table(
        doc,
        ["弹幕模式", "面向用户", "展示策略", "产品价值"],
        [
            ["沉浸", "不喜欢弹幕或想认真看剧情的用户。", "关闭或极低密度，只保留必要高光互动。", "保证短剧内容本身不被干扰。"],
            ["轻聊", "默认用户。", "低密度、半透明、避开关键字幕和人脸区域。", "有陪伴感，但不抢画面。"],
            ["狂欢", "弹幕爱好者和演示场景。", "更高密度，允许适度重复情绪弹幕。", "强化短剧高光时的集体情绪爆发。"],
        ],
        [1.15, 1.75, 2.2, 1.4],
    )
    add_paragraph(
        doc,
        "最终选择这种弹幕方案，是因为短剧弹幕的核心矛盾不是“显示或不显示”，而是“什么时候显示、显示多少、显示给谁、是否会剧透”。三档模式和七层治理能同时满足观看体验、社区氛围和内容安全要求。",
    )

    doc.add_heading("2.5.2 声音模块：从用户授权样本到剧情带入语音", level=3)
    add_paragraph(
        doc,
        "声音模块的目标不是做一个通用语音聊天机器人，而是在片尾二创、陪看小助手和剧情卡中提供更强的代入感。用户上传或录入授权声音样本后，系统只为预设台词生成音频，并缓存结果。这样既能展示“用我的声音进入剧情”的创新点，也能控制隐私、成本和生成风险。",
    )
    add_table(
        doc,
        ["问题", "如果直接实时生成会怎样", "当前解决方式"],
        [
            ["等待时间", "用户点击后等待模型生成，容易打断片尾二创节奏。", "预设文本提前生成或首次生成后缓存 mp3。"],
            ["声音隐私", "如果样本路径公开或被任意调用，会带来隐私风险。", "后端保存 voice profile，前端只拿生成音频播放地址。"],
            ["滥用风险", "开放任意文本会导致冒充、低俗、诈骗等风险。", "只支持固定台词和受控场景，不开放自由文本实时合成。"],
            ["部署成本", "把模型放到手机端会增加包体和设备压力。", "声音生成作为服务端能力，客户端不携带模型文件。"],
            ["重复生成", "同一句话多次生成浪费时间和算力。", "按用户、声音 profile、文本 hash 做缓存。"],
        ],
        [1.2, 2.6, 2.7],
    )
    add_code_block(
        doc,
        [
            "用户录音 / 上传样本",
            "  -> 用户确认授权文本：同意利用录入声音生成音频",
            "  -> 后端保存 voice profile",
            "  -> 选择预设台词，例如片尾二创分镜旁白、小半陪看提示",
            "  -> 调用声音服务生成 mp3",
            "  -> 写入 voice_cache",
            "  -> 前端在二创图片或剧情卡点击时播放",
        ],
    )
    add_table(
        doc,
        ["声音类型", "用途", "生成方式", "播放场景"],
        [
            ["原版声音", "模拟剧中角色或固定旁白风格。", "选取剧中合适参考音频或预设音色，生成固定台词。", "北往片尾二创、剧情卡旁白。"],
            ["用户声音带入版", "让用户感觉自己参与剧情。", "使用用户授权声音样本生成同一批预设台词。", "片尾二创选择分支后播放。"],
            ["小半陪看声音", "作为固定陪看角色提示用户。", "使用固定音色生成轻提示和战报文案。", "高光前后提示、片尾观看战报。"],
        ],
        [1.35, 1.8, 1.95, 1.4],
    )
    add_table(
        doc,
        ["候选方案", "优点", "风险/不足", "最终判断"],
        [
            ["浏览器端 TTS", "实现简单，延迟低。", "声音个性化弱，不能体现用户声音带入。", "只适合兜底，不作为亮点。"],
            ["手机端部署声音模型", "离线可用，不依赖服务端。", "模型文件大，部署复杂，隐私和性能压力高。", "当前不采用。"],
            ["服务端实时生成", "能力完整，用户输入自由。", "等待时间和滥用风险高。", "不作为主链路。"],
            ["服务端预设文本生成 + 缓存", "稳定、可审核、可控成本。", "灵活性低于实时生成。", "当前采用。"],
        ],
        [1.25, 1.75, 2.15, 1.35],
    )
    add_paragraph(
        doc,
        "声音模块的最终取舍是“牺牲一部分实时自由度，换取稳定体验和安全边界”。对于比赛展示和短剧观看场景，用户不需要任意说什么都生成，反而更需要在关键剧情节点听到贴合剧情的几句声音。因此固定文本、预生成、缓存播放更适合当前产品阶段。",
    )
    add_table(
        doc,
        ["安全要求", "当前说明", "后续必须补齐"],
        [
            ["授权", "录入声音前显示明确授权文本。", "正式声音授权协议和撤回授权入口。"],
            ["删除", "当前保留 voice profile 管理思路。", "删除声音样本并清理对应缓存音频。"],
            ["权限", "前端只播放音频，不直接暴露样本路径。", "对象存储私有桶和签名 URL。"],
            ["审计", "生成记录可追踪。", "管理员审计但不随意下载用户原始声音。"],
        ],
        [1.25, 2.05, 3.2],
    )
    add_paragraph(
        doc,
        "这部分是项目区别于普通短剧播放器的重要亮点：普通播放器最多让用户发弹幕，而半句可以让用户在片尾拓展剧情中用自己的声音听到下一段故事。它既有 AIGC 的展示效果，也和用户成长、片尾二创、小半陪看形成了连贯体验。",
    )

    doc.add_heading("2.5.3 两个模块如何形成联动", level=3)
    add_table(
        doc,
        ["联动点", "弹幕模块贡献", "声音模块贡献", "形成的体验"],
        [
            ["高光前", "弹幕减少剧透，轻聊模式保留少量情绪铺垫。", "小半可用轻提示提醒用户将进入关键剧情。", "用户既不被剧透，又能感到陪伴。"],
            ["高光中", "狂欢模式制造集体情绪，弹幕可点赞回复。", "声音不抢主剧情，只在互动或二创节点出现。", "视频主体仍是中心，互动层服务剧情。"],
            ["片尾后", "用户可把感受和二创结果分享到动态。", "用户声音带入版可作为可分享资产。", "观看结束后仍有延展内容和社交传播。"],
        ],
        [1.25, 1.85, 1.85, 1.55],
    )
    add_paragraph(
        doc,
        "因此，弹幕和声音不是两个孤立功能。弹幕解决“多人共同观看的氛围”，声音解决“用户个人带入剧情的代入感”。两者一起支撑了半句的产品定位：短剧陪伴，而不是普通播放。",
    )


def add_technical_doc(doc: Document) -> None:
    doc.add_heading("二、项目技术文档", level=1)
    doc.add_heading("2.1 模块分析和拆解", level=2)
    add_code_block(
        doc,
        [
            "短剧素材/封面/字幕",
            "  -> 导入脚本初始化剧集、视频地址和演示数据",
            "  -> 大模型离线生成高光表达、贴图策略、弹幕复审候选和二创文案",
            "  -> 人工复核写回数据库",
            "  -> FastAPI 下发剧集详情、高光时间轴、弹幕、二创和声音资产",
            "  -> Web 播放页按时间触发互动、贴图、小半提示和片尾二创",
            "  -> 用户点击/选择/弹幕/同看数据上报",
            "  -> 后台统计和复核页继续优化策略",
        ],
    )
    add_table(
        doc,
        ["模块", "职责", "关键数据/接口"],
        [
            ["内容导入", "扫描视频库、建立短剧和剧集数据、绑定封面和视频地址。", "dramas、episodes、media URL、cover URL"],
            ["高光与体验配置", "存储高光时间、类型、情绪、按钮、贴图窗、来源、模型版本和复核状态。", "highlights、episode_experience_configs"],
            ["播放客户端", "展示短剧、播放视频、监听进度、触发互动、上报行为、进入片尾二创。", "/api/dramas、/api/episodes/{id}、互动上报接口"],
            ["弹幕系统", "导入弹幕、分模式展示、点击点赞回复、治理剧透和低俗内容。", "/api/episodes/{id}/danmaku、governance summary"],
            ["片尾二创", "管理分支、分镜图片、选择节点、语音播放和用户选择记录。", "remix configs、image/audio cache"],
            ["声音资产", "保存授权声音样本，按固定文本生成并缓存 mp3。", "voice profile、voice cache"],
            ["社交同看", "好友申请、聊天、房间成员、同看动态、称号徽章展示。", "friends、chat、rooms、badges"],
            ["复核后台", "供人工编辑高光、贴图窗、弹幕治理、二创精选和模型来源。", "admin/review APIs"],
        ],
        [1.25, 3.2, 2.05],
    )

    doc.add_heading("2.2 整体流程图说明", level=2)
    add_code_block(
        doc,
        [
            "[内容侧]",
            "视频/字幕/剧情备注 -> LLM 离线理解 -> 高光/贴图/二创草稿 -> 人工复核 -> 数据库",
            "",
            "[服务侧]",
            "数据库 + 本地资产 -> FastAPI -> 客户端接口 / 复核接口 / 统计接口",
            "",
            "[端上体验]",
            "选片首页 -> 播放器 -> 时间轴触发高光 -> 弹幕互动 -> 片尾二创 -> 战报/成长/社交",
            "",
            "[反馈闭环]",
            "点击、选择、弹幕、同看、二创选择 -> 服务端上报 -> 统计 -> 复核和模型迭代",
        ],
    )

    doc.add_heading("2.3 核心模块技术选型和理由", level=2)
    add_table(
        doc,
        ["模块", "当前选型", "选择理由", "后续演进"],
        [
            ["前端/客户端", "电脑端 Web 主线 + Android 原生支线", "Web 修改快、演示稳定、和服务端联调成本低；Android 支线用于验证未来 App 能力。", "Web 继续作为交付主线；Android 逐步复用接口追赶体验，iOS 后续再评估。"],
            ["服务端", "FastAPI", "接口开发快，适合 AI 工具链和 Python 数据处理脚本集成；API 文档自动生成。", "生产部署时保留 FastAPI，增加 Nginx、进程守护和监控。"],
            ["数据库", "SQLite", "当前演示阶段轻量、无需额外服务、便于备份和回溯。", "正式上线迁移 PostgreSQL。"],
            ["文件存储", "本地文件系统", "适合本地视频、贴图、二创图片和音频缓存快速调试。", "公网访问时迁移对象存储，数据库保存对象 key。"],
            ["大模型", "离线调用大模型/多模态模型", "适合高光理解、文案、贴图策略和二创提示词，不放入实时播放主链路。", "保留人工复核和版本字段，逐步沉淀训练数据。"],
            ["小模型", "弹幕治理小模型雏形", "用于低成本判断弹幕风险、模式和置信度，减少人工审核量。", "后续再训练高光二分类、类型分类和模板推荐模型。"],
            ["声音服务", "本地 CosyVoice 类服务 + mp3 缓存", "支持 3 秒声音样本复刻和固定台词生成，适合片尾二创。", "生产环境作为服务端能力，不放入手机客户端。"],
        ],
        [1.15, 1.35, 2.55, 1.45],
    )
    add_tech_selection_expansion(doc)

    doc.add_heading("2.4 大模型 / AI 能力使用说明", level=2)
    add_paragraph(
        doc,
        "项目中的 AI 能力采用“离线生成、缓存、人工复核、服务端下发”的策略。这样可以避免用户观看时等待模型，降低调用成本，也减少模型误判直接影响体验的风险。",
    )
    add_table(
        doc,
        ["AI 能力", "输入", "输出", "当前落地方式"],
        [
            ["高光表达升级", "剧集信息、字幕/剧情摘要、已有高光时间锚点、人工备注。", "高光名称、情绪、互动按钮、贴图建议、解释理由。", "20 集已完成一轮 LLM 草稿升级，标记为 llm_draft。"],
            ["贴图策略生成", "高光类型、台词关键词、剧情情绪、主题风格。", "贴图文案、出现时间窗、动效建议、点击反馈。", "写入体验配置，复核后进入播放页。"],
            ["弹幕治理候选", "弹幕文本、出现时间、剧情揭晓点、历史审核结果。", "是否剧透、低俗、不相关、适合模式、是否进入复审队列。", "七层治理中的语义和 LLM 候选层。"],
            ["片尾 AI 二创", "北往第一集剧情、主角形象参考、三条分支方向、个性化选项。", "分支文案、分镜台词、图片生成提示词、语音台词。", "采用预生成图片和音频缓存，不做实时视频生成。"],
            ["声音生成", "授权声音样本、固定台词文本、语速参数。", "mp3 音频文件。", "原版声音和用户声音带入版分开缓存。"],
            ["陪看战报", "用户本集互动、选择、点击、奖励和观看上下文。", "小半陪看提示和片尾战报文案。", "前端展示固定 AI 陪看员，不做开放式智能体。"],
        ],
        [1.35, 1.65, 1.8, 1.7],
    )
    add_feature_deep_dive(doc)

    doc.add_heading("2.6 Prompt 设计与组织", level=2)
    add_bullets(
        doc,
        [
            "角色约束：要求模型以短剧互动产品策划和剧情理解助手身份工作，避免输出泛泛影评。",
            "输入结构：按剧集信息、剧情摘要、时间轴、字幕证据、已有高光、用户要求分块输入。",
            "输出结构：要求 JSON 或固定字段输出，包括 highlight_type、emotion、title、buttons、sticker_windows、reason、model_version。",
            "质量约束：不能擅自改变已确认高光时间；不能剧透未到达时间点；每集高光保持稀疏，避免全程都是弹层。",
            "复核约束：模型结果标记为 llm_draft，进入人工复核后才作为稳定展示内容。",
        ]
    )

    doc.add_heading("2.7 工程难点与解决方案", level=2)
    add_table(
        doc,
        ["难点", "具体问题", "解决方案"],
        [
            [
                "观看页叠加内容多，容易遮挡剧情",
                "高光弹层、弹幕、贴图、片尾二创、全屏控制栏都在视频上方，若设计不当会影响观看。",
                "设置视频安全区；控制栏鼠标移动才出现；高光弹层可关闭；二创入口延后到片尾；二创进入后采用独立全屏覆盖层。",
            ],
            [
                "AI 生成不稳定，实时调用成本高",
                "图片、声音、剧情文案和高光理解如果实时生成，延迟、失败率和成本都会影响演示。",
                "采用离线生成 + 缓存 + 人工复核；片尾二创使用预生成图片和 mp3；服务端只下发稳定资产地址。",
            ],
            [
                "弹幕既要热闹又不能剧透或低俗",
                "真实弹幕量大，人工逐条审核成本高；只看文本无法判断是否剧透。",
                "建立七层治理：规则、时间感知、语义分类、聚类去重、小模型、大模型复审候选、人工复核，既降低成本又保留可解释性。",
            ],
        ],
        [1.45, 2.2, 2.85],
    )

    doc.add_heading("2.8 部署与上线规划", level=2)
    add_paragraph(
        doc,
        "当前不启动真实部署，规划目标是为未来上线保留清晰路线。短期继续以本地 Web 和临时公网演示为主，Android 支线只做原生能力验证；当重点剧集体验、数据结构、内容安全和声音授权稳定后，再进入准生产迁移。",
    )
    add_table(
        doc,
        ["阶段", "目标", "主要任务", "验收标准"],
        [
            ["A 展示稳定期", "保证电脑端 Web 可稳定演示。", "打磨播放页、高光弹层、贴图动效、二创入口、小半战报。", "北往第一集完整流程稳定可演示。"],
            ["B 内容规模化期", "从样板集扩展到更多剧集。", "复核 20 集 LLM 草稿，建立每部剧主题播放器和贴图策略。", "至少 4 部剧风格差异明显。"],
            ["C 治理产品化期", "让复核后台真正可操作。", "弹幕七层结果可视化，高光、贴图、二创精选可筛选复核。", "能解释每条内容的来源和处理结果。"],
            ["D 受控公网期", "少量外部用户可访问。", "临时公网或测试域名，限制账号和内容发布。", "外部用户能完成核心体验。"],
            ["E 准生产期", "具备长期运行条件。", "PostgreSQL、对象存储、HTTPS、备份、监控、回滚。", "数据可备份恢复，资产可对账。"],
        ],
        [1.15, 1.55, 2.65, 1.15],
    )

    doc.add_heading("2.9 工作项拆分与排期", level=2)
    add_table(
        doc,
        ["角色/方向", "职责", "近期工作", "里程碑"],
        [
            ["产品设计", "定义观看体验、高光类型、二创分支、社交规则和演示路径。", "继续确认重点剧集高光、二创脚本和展示优先级。", "形成 2-4 部样板剧完整体验。"],
            ["前端/客户端", "实现 Web 首页、播放页、弹幕、贴图、二创、同看、我的等界面，并保留 Android 原生支线。", "继续打磨电脑端观看沉浸感；Android 先复用接口完成最小播放闭环。", "Web 播放页可稳定录屏展示；Android 支线不影响主线交付。"],
            ["后端", "提供内容、互动、弹幕、社交、声音、二创和复核接口。", "完善复核和统计接口，保留 PostgreSQL 迁移准备。", "接口文档和数据结构稳定。"],
            ["AI/模型", "负责 LLM 标注、贴图策略、弹幕治理、小模型和声音生成。", "复核 20 集 LLM 草稿，完善弹幕治理解释。", "模型结果可解释、可追踪。"],
            ["数据/内容运营", "整理视频、封面、弹幕、剧集摘要和二创素材。", "重点剧集人工抽检，确认高光时间和素材质量。", "演示数据可靠无明显错配。"],
            ["部署/文档", "维护运行说明、技术文档、交付材料和未来上线方案。", "保持 PROJECT_STATUS、RUN_LOG 和交付文档最新。", "答辩材料完整可交付。"],
        ],
        [1.25, 1.65, 2.45, 1.15],
    )
    add_paragraph(
        doc,
        "建议排期：第 1 周完成重点剧集观看体验稳定；第 2 周完成 20 集内容抽检和弹幕治理展示；第 3 周完成二创、声音和小半战报样板；第 4 周完成文档、录屏脚本、演示检查和答辩材料。若后续进入正式上线，再单独安排 PostgreSQL、对象存储、HTTPS、备份和监控迁移周期。",
    )

    doc.add_heading("2.10 交付文档和演示材料体系", level=2)
    add_paragraph(
        doc,
        "为了让项目可以被评委、老师、队友或后续 AI 快速接手，当前仓库已经把运行、接口、数据库、模型、隐私、安全、生产计划和演示流程拆成独立文档。这样做的目的不是堆文档，而是降低答辩和后续迭代时的解释成本。",
    )
    add_table(
        doc,
        ["材料", "用途", "当前状态"],
        [
            ["README.md", "项目入口，说明当前 V1 范围、本地运行方式和文档索引。", "已同步当前主线。"],
            ["PROJECT_DELIVERY_MANUAL.md", "比赛提交和老师快速理解项目的总说明。", "已完成第一版。"],
            ["PRESENTATION_SCRIPT.md", "录屏脚本、答辩讲解稿和常见问答。", "已完成第一版。"],
            ["DEMO_CHECKLIST.md", "演示前逐项检查服务、素材、播放、高光、弹幕、二创、同看和复核页。", "已完成第一版。"],
            ["API_REFERENCE.md / DATABASE_SCHEMA.md", "说明服务端接口、核心表、数据关系和后续数据库演进。", "已完成第一版。"],
            ["MODEL_USAGE.md", "说明大模型、小模型预留、弹幕治理、二创和声音资产链路。", "已完成第一版。"],
            ["architecture.md", "说明当前 Web 主线、七层弹幕治理、小半陪看、二创和声音缓存架构。", "已更新到当前真实状态。"],
            ["PRIVACY_AND_CONTENT_SAFETY.md", "说明声音授权、头像照片、弹幕社交治理、AI 二创声明和上线前安全要求。", "已完成第一版。"],
            ["PRODUCTION_PLAN.md", "说明未来云服务器、PostgreSQL、对象存储、HTTPS、备份、监控和迁移路线。", "已完成第一版。"],
        ],
        [2.0, 3.35, 1.15],
    )


def add_summary(doc: Document) -> None:
    doc.add_heading("三、项目总结和自评", level=1)
    doc.add_heading("3.1 对照课题总结", level=2)
    add_paragraph(
        doc,
        "本项目围绕“基于短剧剧情理解的即时互动激发系统”完成了从内容导入、剧情高光理解、互动配置、端上渲染、用户行为上报、弹幕治理、片尾 AI 二创到复核统计的完整闭环。当前版本已经不只是单一播放器 Demo，而是具备真实服务端、数据结构、后台复核、Web 客户端、AI 资产生成和产品化规划的全栈 MVP。",
    )
    add_table(
        doc,
        ["课题要求", "当前完成情况", "自评"],
        [
            ["短剧内容理解", "已支持高光时间轴、类型、情绪、文案、来源、模型版本和人工复核。", "达到 MVP 要求，后续可进一步训练高光识别小模型。"],
            ["AIGC 内容生成", "已落地片尾二创图片分镜、剧情卡、声音生成缓存、小半陪看战报和贴图策略。", "具备展示亮点，视频生成暂不作为当前主线。"],
            ["短剧端上渲染", "Web 播放页支持高光弹层、弹幕、贴图动效、全屏、二创覆盖层。", "电脑端展示较完整，移动原生仍需后续探索。"],
            ["服务端闭环", "FastAPI + SQLite 管理内容、互动、弹幕、社交、声音和复核接口。", "演示闭环完整，生产化需迁移 PostgreSQL 和对象存储。"],
            ["可迭代性", "已有状态记录、运行日志、API、数据库、部署、模型、隐私和生产计划文档。", "便于后续团队或 AI 接手。"],
        ],
        [1.55, 3.25, 1.7],
    )

    doc.add_heading("3.2 项目收获", level=2)
    add_bullets(
        doc,
        [
            "产品层面：高光互动不能只追求弹窗数量，必须保持剧情对比度和低打扰，否则会破坏短剧观看体验。",
            "技术层面：AI 能力最适合放在离线生产和复核链路中，端上只加载稳定结果，用户体验更可靠。",
            "工程层面：复杂 AI 产品必须保留 source、model_version、confidence、review_status 等字段，否则后续无法解释和优化。",
            "体验层面：弹幕、贴图、二创和陪看都需要服务剧情，而不是为了炫技堆效果。",
            "治理层面：弹幕审核需要结合时间轴，单纯文本审核无法解决剧透问题。",
        ]
    )

    doc.add_heading("3.3 不足之处", level=2)
    add_bullets(
        doc,
        [
            "高光识别小模型尚未完整训练和部署，目前高光主要依赖大模型草稿和人工复核。",
            "当前生产部署仍停留在规划层面，真实 PostgreSQL、对象存储、HTTPS、备份和监控尚未执行。",
            "片尾 AI 二创采用图片分镜和语音缓存，真实视频生成效果和成本还不稳定，暂未作为正式展示主线。",
            "声音资产已经形成产品雏形，但正式开放前还需要更完整的授权、删除、权限和审计机制。",
            "Android 原生迁移已经作为支线探索，但当前效果还未达到 Web 主线展示质量，因此当前交付主线仍保持电脑端 Web；未来支线继续围绕播放、全屏、权限和覆盖层逐步补齐。",
            "20 集 LLM 升级结果仍需要继续人工抽检，重点剧集体验优先于全量复杂特效覆盖。",
        ]
    )

    doc.add_heading("3.4 后续迭代方向", level=2)
    add_numbers(
        doc,
        [
            "继续打磨电脑端 Web 观看体验，优先把北往和那年冬至做成稳定样板。",
            "复核 20 集 LLM 高光和贴图草稿，建立不同题材短剧的主题播放器和互动风格。",
            "完善弹幕治理后台，让规则、时间感知、语义、小模型、LLM 候选和人工复核结果都可解释。",
            "将声音资产、片尾二创和小半陪看统一成可配置的内容生产流程。",
            "在项目稳定后，再按生产计划迁移 PostgreSQL、对象存储、HTTPS、备份和监控。",
            "重新评估 Android/iOS 客户端迁移，优先保证端上播放、全屏、语音权限和互动覆盖层质量。",
        ]
    )

    add_callout(
        doc,
        "自评结论",
        "当前版本已经完成比赛展示所需的核心闭环，并在 AI 内容理解、弹幕治理、片尾二创、声音带入和陪看战报方面形成明显亮点。不足主要集中在生产部署、小模型规模化训练、隐私合规完善和原生客户端质量上。后续应继续坚持“先稳定体验，再扩展能力”的路线。",
    )

    doc.add_heading("3.5 当前可交付状态与风险控制", level=2)
    add_table(
        doc,
        ["维度", "当前状态", "风险控制"],
        [
            ["比赛展示", "电脑端 Web 主线已经具备首页、播放、高光、弹幕、二创、声音、同看、个人主页和复核页。", "优先使用北往第一集等重点样板演示，避免现场切到未复核剧集。"],
            ["技术讲解", "已有架构、接口、数据库、模型、隐私安全、生产计划和演示脚本文档。", "答辩时按“内容理解、端上渲染、AI 生成、治理闭环”四条线讲解。"],
            ["AI 生成", "采用离线生成、缓存、人工复核，不依赖现场实时生成。", "模型失败不影响核心播放链路。"],
            ["用户数据", "当前主要是演示数据和本地数据。", "正式上线前必须补齐授权、删除、举报、审核和日志脱敏。"],
        ],
        [1.35, 3.1, 2.05],
    )


def build() -> None:
    doc = setup_document()
    add_cover(doc)
    add_overview(doc)
    doc.add_page_break()
    add_technical_doc(doc)
    doc.add_page_break()
    add_summary(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
